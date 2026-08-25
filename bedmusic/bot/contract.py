"""Builds the signed contract as a .docx.

The wording follows the template supplied for this project (отчуждение
исключительного права на инструментальную композицию). Section 3 is the one
deliberate departure: the template settles in roubles by bank transfer or cash,
while deals here settle in crypto through the bot's escrow, so that section
states the escrow mechanics instead.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from . import money

BLANK = "____________________"

MONTHS_GENITIVE = [
    "января", "февраля", "марта", "апреля", "мая", "июня",
    "июля", "августа", "сентября", "октября", "ноября", "декабря",
]


def _fill(value: Optional[str], width: int = 20) -> str:
    """A filled value, or an underscore run to be completed by hand."""
    return value if value else "_" * width


def _long_date(day: date) -> str:
    """«24» августа 2026 г. — the form a Russian contract uses."""
    return f"«{day.day:02d}» {MONTHS_GENITIVE[day.month - 1]} {day.year} г."


class _Builder:
    def __init__(self) -> None:
        self.doc = Document()
        style = self.doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        for section in self.doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(1.5)

    def title(self, text: str, size: int = 13) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)

    def para(self, text: str = "", bold: bool = False, align: str = "justify") -> None:
        p = self.doc.add_paragraph()
        p.alignment = {
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }[align]
        run = p.add_run(text)
        run.bold = bold

    def heading(self, text: str) -> None:
        self.para("")
        self.para(text, bold=True, align="left")

    def signature_block(self, caption: str, image: Optional[Path], name: str) -> None:
        """Signature image if we have one, otherwise a line to sign by hand."""
        if image and image.exists():
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            try:
                p.add_run().add_picture(str(image), width=Cm(4))
            except Exception:  # noqa: BLE001 — unreadable image must not lose the deal
                p.add_run("__________")
            self.para(f"{caption} / {name} /", align="left")
        else:
            self.para(f"{caption}: __________ / {_fill(name)} /", align="left")

    def save(self, path: Path) -> Path:
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))
        return path


def _duration(seconds: int) -> str:
    return f"{seconds // 60} мин {seconds % 60:02d} сек" if seconds else "______ мин ______ сек"


def render(
    *,
    path: Path,
    title: str,
    price_units: int,
    price_currency: str,
    fields: dict,
    duration_seconds: int = 0,
    seller_signature: Optional[Path] = None,
    buyer_signature: Optional[Path] = None,
    contract_date: Optional[date] = None,
    deal_id: Optional[int] = None,
    material_names: Optional[list[str]] = None,
) -> Path:
    """Render the contract. Missing answers become blanks, never crash."""
    f = fields
    today = contract_date or date.today()
    amount = money.format_amount(price_units, price_currency)
    exact = money.amount_exact(price_units, price_currency)

    b = _Builder()
    b.title("ДОГОВОР")
    b.title("об отчуждении исключительного права", 11)
    b.title("на музыкальное произведение (инструментальную композицию)", 11)
    b.para("")

    p = b.doc.add_paragraph()
    p.add_run(f"г. {_fill(f.get('city'))}")
    p.add_run("\t\t\t\t\t")
    p.add_run(_long_date(today))
    b.para("")

    if deal_id:
        b.para(f"Сделка № {deal_id} в сервисе Bed Music", align="left")
        b.para("")

    b.para(
        f"Гражданин(ка) {_fill(f.get('seller_fio'), 40)}, дата рождения "
        f"{_fill(f.get('seller_birthdate'), 12)}, паспорт серия и номер "
        f"{_fill(f.get('seller_passport'), 14)}, выдан {_fill(f.get('seller_passport_issued'), 40)}, "
        f"зарегистрирован(а) по адресу: {_fill(f.get('seller_address'), 40)}, именуемый(ая) "
        "в дальнейшем «Автор» (Правообладатель), с одной стороны,"
    )
    b.para(
        f"и гражданин(ка) {_fill(f.get('buyer_fio'), 40)}, дата рождения "
        f"{_fill(f.get('buyer_birthdate'), 12)}, паспорт серия и номер "
        f"{_fill(f.get('buyer_passport'), 14)}, выдан {_fill(f.get('buyer_passport_issued'), 40)}, "
        f"зарегистрирован(а) по адресу: {_fill(f.get('buyer_address'), 40)}, именуемый(ая) "
        "в дальнейшем «Приобретатель», с другой стороны,"
    )
    b.para("совместно именуемые «Стороны», заключили настоящий Договор о нижеследующем.")

    # 1
    b.heading("1. ПРЕДМЕТ ДОГОВОРА")
    b.para(
        "1.1. Автор передаёт Приобретателю в полном объёме исключительное право на созданное "
        "творческим трудом Автора музыкальное произведение — инструментальную композицию "
        "(далее — «Произведение», «Бит»), а Приобретатель обязуется уплатить Автору "
        "предусмотренное настоящим Договором вознаграждение."
    )
    b.para("1.2. Характеристики Произведения:")
    b.para(f"— Название: {_fill(title, 40)}")
    b.para(f"— Продолжительность: {_duration(duration_seconds)}")
    b.para(f"— Темп (BPM): {_fill(f.get('beat_bpm'), 6)} ; тональность: {_fill(f.get('beat_key'), 10)}")
    b.para(f"— Дата создания: {_fill(f.get('beat_created'), 12)}")
    b.para("1.3. Автор передаёт Приобретателю следующие материалы (далее — «Материалы»):")
    b.para(f"— мастер-файл в формате WAV ({_fill(f.get('wav_quality'), 20)});")
    b.para("— файл в формате MP3 (320 кбит/с);")
    b.para(f"— раздельные дорожки (stems) — {_fill(f.get('has_stems'), 6)};")
    b.para(f"— проектный файл (при наличии) — {_fill(f.get('has_project'), 6)}.")
    b.para(
        "1.4. Материалы передаются Приобретателю в электронном виде через сервис Bed Music "
        "в момент подписания настоящего Договора обеими Сторонами, одновременно с "
        "перечислением вознаграждения Автору. Материалы, которые не могут быть переданы "
        "через сервис по техническим причинам, направляются на адрес электронной почты "
        f"Приобретателя {_fill(f.get('buyer_email'), 30)} в срок не позднее "
        f"{_fill(f.get('delivery_days'), 4)} календарных дней с момента подписания."
    )
    b.para(
        "1.5. Исключительное право переходит к Приобретателю в момент подписания Сторонами "
        "настоящего Договора либо в момент полной оплаты вознаграждения — в зависимости от того, "
        "что наступит позднее."
    )

    # 2
    b.heading("2. ОБЪЁМ ПЕРЕДАВАЕМЫХ ПРАВ")
    b.para(
        "2.1. Исключительное право передаётся в полном объёме, без ограничения по территории "
        "(весь мир) и на весь срок действия исключительного права."
    )
    b.para(
        "2.2. Приобретатель вправе использовать Произведение любым не противоречащим закону "
        "способом, в том числе: воспроизводить, распространять, публично исполнять, доводить до "
        "всеобщего сведения, перерабатывать (в том числе накладывать вокал, изменять аранжировку, "
        "сводить и мастерить), включать в состав аудиовизуальных произведений, размещать на "
        "цифровых площадках (стриминговые сервисы, видеохостинги, социальные сети), а также "
        "передавать полученные права третьим лицам."
    )
    b.para(
        "2.3. С момента перехода исключительного права Автор не вправе самостоятельно использовать "
        "Произведение указанными способами, а также отчуждать права на него третьим лицам, в том "
        "числе продавать Бит другим покупателям."
    )
    b.para(
        "2.4. Личные неимущественные права (право авторства, право на имя, право на "
        "неприкосновенность произведения в части, установленной законом) сохраняются за Автором "
        "и не отчуждаются."
    )
    b.para(
        "2.5. При использовании Произведения Приобретатель обязуется указывать авторство Бита "
        f"в форме: «prod. by {_fill(f.get('seller_alias'), 20)}» — в описании релиза, метаданных "
        "трека и на цифровых площадках, где это технически возможно. Стороны вправе согласовать "
        "иную форму указания либо отказ от указания, оформив это дополнительным соглашением."
    )

    # 3 — rewritten for crypto escrow
    b.heading("3. ВОЗНАГРАЖДЕНИЕ И ПОРЯДОК РАСЧЁТОВ")
    b.para(f"3.1. Размер вознаграждения Автора составляет {amount} ({exact}).")
    b.para(
        "3.2. Расчёты производятся в криптовалюте через сервис Bed Music, выступающий "
        "эскроу-агентом Сторон. Сумма вознаграждения блокируется на счёте Приобретателя в "
        "сервисе в момент направления запроса на покупку и перечисляется Автору после подписания "
        "настоящего Договора обеими Сторонами."
    )
    b.para(
        "3.3. Адрес кошелька Автора для получения вознаграждения: "
        f"{_fill(f.get('seller_payout'), 40)}"
    )
    b.para(
        "3.4. Обязанность Приобретателя по оплате считается исполненной с момента блокирования "
        "суммы вознаграждения эскроу-агентом. Расчёты между Сторонами считаются завершёнными "
        "с момента подтверждения перевода в соответствующей сети."
    )
    b.para(
        "3.5. Если Договор не подписан обеими Сторонами либо сделка отменена, заблокированная "
        "сумма возвращается Приобретателю в полном объёме."
    )
    b.para(
        "3.6. Стороны осознают, что курс криптовалюты может изменяться, и принимают на себя "
        "связанные с этим риски. Вознаграждение является окончательным; роялти от использования "
        "Произведения настоящим Договором не предусмотрены, если иное не согласовано письменно."
    )

    # 4
    b.heading("4. ГАРАНТИИ И ЗАВЕРЕНИЯ АВТОРА")
    b.para(
        "4.1. Автор гарантирует, что является единственным автором Произведения и обладателем "
        "исключительного права на него в полном объёме."
    )
    b.para(
        "4.2. Автор гарантирует, что Произведение создано им лично, ранее не отчуждалось третьим "
        "лицам, не находится под залогом, арестом, не является предметом действующих лицензионных "
        "договоров и не размещено на площадках продажи битов на условиях эксклюзива."
    )
    b.para(
        "4.3. Автор гарантирует, что Произведение не содержит фрагментов (сэмплов, лупов, "
        "вокальных партий) чужих произведений, права на которые принадлежат третьим лицам, "
        "за исключением следующих элементов, использованных на законных основаниях:"
    )
    b.para(f.get("samples") or "отсутствуют")
    b.para(
        "4.4. В случае предъявления третьими лицами претензий, связанных с нарушением их прав "
        "содержанием Произведения, Автор обязуется урегулировать такие претензии своими силами "
        "и за свой счёт, а также возместить Приобретателю документально подтверждённые убытки."
    )

    # 5
    b.heading("5. ОТВЕТСТВЕННОСТЬ СТОРОН")
    b.para(
        "5.1. За нарушение сроков передачи Материалов Автор уплачивает Приобретателю пеню в "
        "размере 0,1 % от суммы вознаграждения за каждый день просрочки, но не более 10 % "
        "от суммы вознаграждения."
    )
    b.para(
        "5.2. За нарушение сроков оплаты Приобретатель уплачивает Автору пеню в размере 0,1 % "
        "от неуплаченной суммы за каждый день просрочки, но не более 10 % от суммы вознаграждения."
    )
    b.para(
        "5.3. В остальном Стороны несут ответственность в соответствии с законодательством "
        "Российской Федерации."
    )

    # 6
    b.heading("6. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ")
    b.para(
        "6.1. Договор вступает в силу с момента подписания его Сторонами и действует до полного "
        "исполнения обязательств."
    )
    b.para(
        "6.2. Все изменения и дополнения к Договору действительны, если совершены в письменной "
        "форме и подписаны Сторонами."
    )
    b.para(
        "6.3. Стороны признают юридическую силу за документами и сообщениями, направленными по "
        "указанным в Договоре адресам электронной почты, а также за экземплярами Договора, "
        "подписанными собственноручно и переданными в виде сканированных копий или фотографий."
    )
    b.para(
        "6.4. Споры разрешаются путём переговоров, а при недостижении согласия — в суде "
        "в соответствии с законодательством Российской Федерации."
    )
    b.para(
        "6.5. Договор составлен в двух экземплярах, имеющих равную юридическую силу, по одному "
        "для каждой из Сторон."
    )
    b.para(
        "6.6. Факт передачи Материалов и отсутствие взаимных претензий подтверждаются Актом "
        "приёма-передачи (Приложение № 1), являющимся неотъемлемой частью Договора."
    )

    # 7 — only when the author is under 18
    if f.get("seller_is_minor") == "ДА":
        b.heading("7. СОГЛАСИЕ ЗАКОННОГО ПРЕДСТАВИТЕЛЯ АВТОРА")
        b.para(
            f"Я, {_fill(f.get('rep_fio'), 40)}, паспорт серия и номер "
            f"{_fill(f.get('rep_passport'), 14)}, выдан {_fill(f.get('rep_passport_issued'), 40)}, "
            "являясь законным представителем (родителем / усыновителем / попечителем) Автора, "
            "ознакомлен(а) с условиями настоящего Договора, выражаю своё согласие на его "
            "заключение и подтверждаю правомерность действий Автора по распоряжению исключительным "
            "правом на Произведение."
        )
        b.para("Подпись законного представителя: ______________ / ______________________________ /")

    # 8
    b.heading("8. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН")
    b.para("")
    b.para("АВТОР", bold=True, align="left")
    b.para(f"ФИО: {_fill(f.get('seller_fio'), 30)}", align="left")
    b.para(f"Дата рождения: {_fill(f.get('seller_birthdate'), 12)}", align="left")
    b.para(f"Паспорт: {_fill(f.get('seller_passport'), 14)}", align="left")
    b.para(f"Выдан: {_fill(f.get('seller_passport_issued'), 30)}", align="left")
    b.para(f"Адрес: {_fill(f.get('seller_address'), 30)}", align="left")
    b.para(f"Телефон: {_fill(f.get('seller_phone'), 16)}", align="left")
    b.para(f"E-mail: {_fill(f.get('seller_email'), 24)}", align="left")
    b.para(f"Псевдоним: {_fill(f.get('seller_alias'), 20)}", align="left")
    b.para(f"Кошелёк для оплаты: {_fill(f.get('seller_payout'), 40)}", align="left")
    b.para("")
    b.signature_block("Подпись", seller_signature, f.get("seller_fio", ""))

    b.para("")
    b.para("ПРИОБРЕТАТЕЛЬ", bold=True, align="left")
    b.para(f"ФИО: {_fill(f.get('buyer_fio'), 30)}", align="left")
    b.para(f"Дата рождения: {_fill(f.get('buyer_birthdate'), 12)}", align="left")
    b.para(f"Паспорт: {_fill(f.get('buyer_passport'), 14)}", align="left")
    b.para(f"Выдан: {_fill(f.get('buyer_passport_issued'), 30)}", align="left")
    b.para(f"Адрес: {_fill(f.get('buyer_address'), 30)}", align="left")
    b.para(f"Телефон: {_fill(f.get('buyer_phone'), 16)}", align="left")
    b.para(f"E-mail: {_fill(f.get('buyer_email'), 24)}", align="left")
    b.para(f"Псевдоним / артист-нейм: {_fill(f.get('buyer_alias'), 20)}", align="left")
    b.para("")
    b.signature_block("Подпись", buyer_signature, f.get("buyer_fio", ""))

    # Приложение № 1
    b.doc.add_page_break()
    b.para(
        f"Приложение № 1 к Договору об отчуждении исключительного права от {_long_date(today)}",
        align="left",
    )
    b.para("")
    b.title("АКТ")
    b.title("приёма-передачи материалов и прав", 11)
    b.para("")
    p = b.doc.add_paragraph()
    p.add_run(f"г. {_fill(f.get('city'))}")
    p.add_run("\t\t\t\t\t")
    p.add_run(_long_date(today))
    b.para("")
    if material_names:
        materials = ", ".join(material_names)
    else:
        materials = "мастер-файл WAV, файл MP3 320 кбит/с"
        if f.get("has_stems") == "ДА":
            materials += ", раздельные дорожки (stems)"
        if f.get("has_project") == "ДА":
            materials += ", проектный файл"
    b.para(
        f"1. Автор передал, а Приобретатель принял материалы Произведения «{_fill(title, 30)}» "
        f"в согласованных форматах, а именно: {materials}."
    )
    b.para(
        "2. Материалы переданы через сервис Bed Music в момент подписания Договора. "
        f"Контактный адрес электронной почты Приобретателя: {_fill(f.get('buyer_email'), 24)}."
    )
    b.para("3. Исключительное право на Произведение перешло к Приобретателю в полном объёме.")
    b.para(f"4. Вознаграждение в размере {amount} ({exact}) получено Автором в полном объёме.")
    b.para(
        "5. Стороны претензий друг к другу не имеют. Обязательства по Договору исполнены "
        "надлежащим образом."
    )
    b.para("")
    b.para("АВТОР", bold=True, align="left")
    b.signature_block("Подпись", seller_signature, f.get("seller_fio", ""))
    b.para("")
    b.para("ПРИОБРЕТАТЕЛЬ", bold=True, align="left")
    b.signature_block("Подпись", buyer_signature, f.get("buyer_fio", ""))

    return b.save(path)
